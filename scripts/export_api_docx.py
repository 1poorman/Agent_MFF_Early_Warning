"""将 design/API_INTERFACE.md 导出为 docx（对齐大赛模板格式）。"""

import pathlib
import sys

from docx import Document
from docx.shared import Pt

ROOT = pathlib.Path(__file__).resolve().parent.parent


def main():
    doc = Document()
    doc.add_heading("中频炉水冷系统预警智能体 · 接入接口文档", 0)
    doc.add_paragraph(
        "依据大赛《工业智能体大赛智能体接入接口文档》规范，"
        "提供 OpenAPI（RESTful）与 MCP 双协议接入。版本 v1.0.0，2026-08-20")

    text = (ROOT / "design" / "API_INTERFACE.md").read_text(encoding="utf-8")
    lines = text.split("\n")
    in_code = False
    for ln in lines[3:]:  # 跳过 md 主标题与引用
        if ln.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(ln)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            continue
        if ln.startswith("## "):
            doc.add_heading(ln[3:], 1)
        elif ln.startswith("### "):
            doc.add_heading(ln[4:], 2)
        elif ln.startswith("#### "):
            doc.add_heading(ln[5:], 3)
        elif ln.strip().startswith(">"):
            continue  # 跳过引用注释
        elif ln.strip():
            doc.add_paragraph(ln)

    out = ROOT / "docs" / "中频炉预警智能体接入接口文档.docx"
    doc.save(str(out))
    print(f"docx saved -> {out}")


if __name__ == "__main__":
    main()
