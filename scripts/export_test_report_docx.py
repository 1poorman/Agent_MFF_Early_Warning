"""将 design/TEST_REPORT.md 导出为 docx（大赛技术报告-测试部分）。"""

import pathlib

from docx import Document
from docx.shared import Pt

ROOT = pathlib.Path(__file__).resolve().parent.parent


def md_to_docx(md_path, out_path, title):
    doc = Document()
    doc.add_heading(title, 0)
    lines = pathlib.Path(md_path).read_text(encoding="utf-8").split("\n")
    in_code = False
    table_buf = []

    def flush_table():
        if not table_buf:
            return
        rows = [r for r in table_buf if not set(r) <= set("|-: ")]
        parsed = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
        if parsed:
            t = doc.add_table(rows=len(parsed), cols=len(parsed[0]))
            t.style = "Light Grid Accent 1"
            for i, row in enumerate(parsed):
                for j, cell in enumerate(row):
                    if j < len(t.rows[i].cells):
                        t.rows[i].cells[j].text = cell
        table_buf.clear()

    for ln in lines[2:]:
        if ln.startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(ln)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            continue
        if ln.strip().startswith("|"):
            table_buf.append(ln)
            continue
        flush_table()
        if ln.startswith("## "):
            doc.add_heading(ln[3:], 1)
        elif ln.startswith("### "):
            doc.add_heading(ln[4:], 2)
        elif ln.startswith("#### "):
            doc.add_heading(ln[5:], 3)
        elif ln.strip().startswith(">"):
            p = doc.add_paragraph(ln.strip().lstrip("> "))
            for r in p.runs:
                r.font.size = Pt(9)
        elif ln.strip():
            doc.add_paragraph(ln)
    flush_table()
    doc.save(str(out_path))
    print(f"docx saved -> {out_path}")


if __name__ == "__main__":
    md_to_docx(ROOT / "design" / "TEST_REPORT.md",
               ROOT / "docs" / "中频炉预警智能体测试报告.docx",
               "中频炉水冷系统预警智能体 · 测试报告")
